# %%
import docx
import re
from typing import List, Tuple, Dict
import os
import json
# import threading
# from transformers import pipeline
import requests
from datetime import datetime

# %%
# MAGIC_PORT = 9399

# %%
def get_tag(content:str):
    content = content.strip('\n').strip()
    
    if content.lower().startswith('phần'):
        return 1
    if content.lower().startswith('chương'):
        return 2
    elif content.lower().startswith('phụ lục'):
        return 1
    elif content.lower().startswith('mục'):
        return 3
    elif content.lower().startswith('điều'):
        return 4
    elif content.lower().startswith('khoản'):
        return 5
    elif content.isupper():
        return 0
    else:
        try:
            re.search(r'^\b([0-9]|[1-9])+[\.|\)|\,]+(\.)?+[A|Ă|Â|B|C|D|Đ|E|Ê|G|H|I|K|L|M|N|O|Ô|Ơ|P|Q|R|S|T|U|Ư|V|X|Y]', 
                      content).group(0)
            return 6
        except:
            pass
        
        try:
            re.search(r'^[a|ă|â|b|c|d|đ|e|ê|g|h|i|k|l|m|n|o|ô|ơ|p|q|r|s|t|u|ư|v|x|y]+([\)|\.|\,])?\s[A|Ă|Â|B|C|D|Đ|E|Ê|G|H|I|K|L|M|N|O|Ô|Ơ|P|Q|R|S|T|U|Ư|V|X|Y]',
                      content).group(0)
            return 7
        except:
            pass
        
        try:
            re.search(r'^\b([0-9]|[1-9])+\.+\b([0-9]|[1-9])+([\.|\)|\,])?\s[A|Ă|Â|B|C|D|Đ|E|Ê|G|H|I|K|L|M|N|O|Ô|Ơ|P|Q|R|S|T|U|Ư|V|X|Y]', 
                      content).group(0)
            return 7
        except:
            pass
        
        try:
            re.search(r'^[a|ă|â|b|c|d|đ|e|ê|g|h|i|k|l|m|n|o|ô|ơ|p|q|r|s|t|u|ư|v|x|y]+\.+\b([0-9]|[1-9])+([\)|\.|\,])?\s[A|Ă|Â|B|C|D|Đ|E|Ê|G|H|I|K|L|M|N|O|Ô|Ơ|P|Q|R|S|T|U|Ư|V|X|Y]', 
                      content).group(0)
            return 8
        except:
            pass
        
        try:
            re.search(r'^\b([0-9]|[1-9])+\.+\b([0-9]|[1-9])+\.+\b([0-9]|[1-9])+([\.|\)|\,])?\s[A|Ă|Â|B|C|D|Đ|E|Ê|G|H|I|K|L|M|N|O|Ô|Ơ|P|Q|R|S|T|U|Ư|V|X|Y]', 
                      content).group(0)
            return 8
        except:
            pass
        
        return 9
        
# %% [markdown]
# ---

# %% [markdown]
# ## Đối với docx

# %%
def extract_docx_layout(document_, 
                 force_close:int=6, 
                 allow_merge:int=7, 
                 sent_separator:str='<\>',
                 meta_header:str='') -> Tuple[List[str], List[str]]:
    document = document_
    if isinstance(document_, docx.document.Document):
        document = [content.text for content in document_.paragraphs]
    elif not isinstance(document_, list):
        return [], []
    
    doc_list = []
    head_list = []

    doc_stack = []
    doc_item = []
    for i in range(len(document)):
        content = document[i]
        text = content.strip('\n').strip().replace('\n', ' ')
        if len(text) == 0: continue
        tag = get_tag(text)
        # print(tag, ':', text)
        if tag == 0 and len(doc_stack) > 0 and doc_stack[-1] in [1,2,3]:
            doc_item[-1] += ' ' + text
            continue
        elif tag == 9 and len(text) > 0 and text[0].islower() and len(doc_stack) > 0:
            doc_item[-1] += ' ' + text
            continue
        elif tag >= allow_merge:
            tag = allow_merge
            
        if len(doc_stack) > 0 and tag == doc_stack[-1] < force_close:
            # print('cut', tag, doc_item, '\n-----------------')
            head_item = doc_item[:-1].copy()
            if meta_header != '':
                head_item.insert(0, meta_header)
            head_list.append(sent_separator.join(head_item))
            _doc = doc_item.pop(-1)
            _tag = doc_stack.pop(-1)
            doc_list.append(_doc)
            
        elif len(doc_stack) > 0 and tag < doc_stack[-1]:
            # print('cut', tag, doc_item[-1])
            pos_h = doc_stack.index(doc_stack[-1]) - 1 
            pos_d = pos_h + doc_stack.count(doc_stack[-1])
            for itag in range(0, doc_stack[-1]):
                try:
                    pos_h = doc_stack.index(itag) 
                except:
                    pass
            # print(pos_h, pos_d, len(doc_stack), tag)
            # print('bef', tag, doc_stack, doc_item[:pos_h+1])
            head_item = doc_item[:pos_h+1]
            if meta_header != '':
                head_item.insert(0, meta_header)
                
            doc_item_ = doc_item[pos_h+1:]
            doc_item_.reverse()
            doc_item = doc_item[:pos_h+1] + doc_item_
            
            while pos_h < pos_d >= 0:
                # print(len(doc_stack), pos_d)
                head_list.append(sent_separator.join(head_item))
                doc_list.append(sent_separator.join(doc_item[pos_d:]))
                # print(head_list[-1])
                # print(doc_list[-1])
                while len(doc_stack) > pos_d:
                    _doc = doc_item.pop(-1)
                    _tag = doc_stack.pop(-1)
                pos_d -= 1
            # print(pos_h, pos_d, len(doc_stack), tag)
            while len(doc_stack) > 0 <= pos_h and doc_stack[pos_h] >= tag:
                _doc = doc_item.pop(-1)
                doc_list[-1] = _doc + sent_separator + doc_list[-1]
                _tag = doc_stack.pop(-1)
                pos_h -= 1
                # print(pos_h, pos_d, len(doc_stack), tag)
            # print('\n-----------------')
            # print('aft', doc_stack, doc_item, '\n-----------------')
        
        if text.lower().startswith('phụ lục'):
            if meta_header != '':
                meta_header_ = sent_separator.join([meta_header] + doc_item + [text])
            else:
                meta_header_ = text
            appendix_head, appendix_doc = extract_docx_layout(document[i+1:], 
                                           force_close=2, 
                                           allow_merge=6, 
                                           sent_separator=sent_separator,
                                           meta_header=meta_header_)
            head_list += appendix_head
            doc_list += appendix_doc
            break 
        
        doc_stack.append(tag)
        doc_item.append(text)
        
        assert len(doc_stack) == len(doc_item), f'Failed when extract document layout!'
    #end-for
    
    # print(head_list)   
    # print(doc_list)
        
    if len(doc_stack) > 0:
        pos = doc_stack.index(doc_stack[-1])
        head_item = doc_item[:pos]
        if meta_header != '':
            head_item.insert(0, meta_header)
        head_list.append(sent_separator.join(head_item))
        doc_list.append(sent_separator.join(doc_item[pos:]))        
        
    try:
        table_content = []
        for table in document_.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        text = para.text.strip('\n').strip().replace('\n', ' ')
                        if text != '':
                            table_content.append(text)
        
        doc_list.append(sent_separator.join(table_content))
        head_list.append('Bảng')
    except:
        pass
    del document_
    del document
    return head_list, doc_list

# %%
# head_list, doc_list = extract_docx_layout(document)

# %%
# i = int(input())
# print(head_list[i])
# print('---')
# print(doc_list[i])

# %% [markdown]
# ---

# %% [markdown]
# ## Đối với PDF

# %% [markdown]
# ### Phân loại PDF: 1. Scanned image / Text-based

# %%
import fitz

def is_vie_content(text):
    try:
        re.search(r'[ă|â|đ|ê|ô|ơ|ư]', text).group(0)
        return True
    except:
        return False
    
def is_scanned_pdf(path, threshold:int=50) -> bool:
    pdf = fitz.open(path)
    for page in pdf:
        # print(page.get_text().strip('\n').strip())
        if page.get_text() is None or \
                len(page.get_text().strip().replace('\n', ' ').split()) < threshold or \
                not is_vie_content(page.get_text().strip()):
            pdf.close()
            return True
    pdf.close()
    return False

# %% [markdown]
# ### Với PDF có thể đọc được

# %%
import textract
# %% [markdown]
# ### Với PDF scanned image

# %%
import pytesseract
import pdf2image
import cv2
import os
import numpy as np
# from transformers import pipeline
# corrector = pipeline("text2text-generation", model="bmd1905/vietnamese-correction", )

# %%
def ocr_extract_text(pdf_path, 
                     dpi:int=350, 
                     thread_count:int=1,
                     margin:Tuple[int, int, int, int]=(200, 200, 150, 50), 
                     config:str='-l vie --psm 6', 
                     im_size:Tuple[int, int]=(2550, 3300),
                     port:int=9400):
    images = pdf2image.convert_from_path(pdf_path, dpi=dpi, thread_count=thread_count, size=im_size)
    document = []
    
    for pil_img in images:
        cv_img = np.array(pil_img).copy()
        _width, _height = pil_img.size
        gray = cv2.cvtColor(cv_img, cv2.COLOR_BGR2GRAY)
        blur = cv2.GaussianBlur(gray, ksize=(9, 9), sigmaX=0)
        
        thresh = cv2.adaptiveThreshold(blur, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY_INV, 11, 30)
        kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (9, 9))
        dilate = cv2.dilate(thresh, kernel, iterations=4)
        contours, _ = cv2.findContours(dilate, mode=cv2.RETR_EXTERNAL, method=cv2.CHAIN_APPROX_SIMPLE)
        
        bbox = []
        for contour in contours:
            x, y, w, h = cv2.boundingRect(contour)
            # print(x, y, w, h)
            if cv2.contourArea(contour) < 10000:
                continue
            if y < margin[0] or y + h > _height - margin[1] \
                or x < margin[2] or x + w > _width - margin[3]: continue
            bbox.insert(0, [x, y, w, h])
            
        # text = ''
        _document = []
        img = np.array(pil_img).copy()
        for x, y, w, h in bbox:
            cropped_image = img[y:y + h, x:x + w]
            _, thresh = cv2.threshold(cropped_image, 127, 255, cv2.THRESH_BINARY)
            text = str(
                pytesseract.image_to_string(thresh, config=config)
                    ).strip()
            # text = re.sub(r'^[\!\?\|\_\*\^\ˆ\"\:\'\/\\\~\`\<\>\&\%\$\#\@\[\]\{\}]', '', text)
            _document.append(text)
        document += requests.post(url=f'http://localhost:{port}/', json={"document": _document}).json().get("generated_text")
    del images
    return document

# %% [markdown]
# ### Process chung

# %%
def extract_pdf_layout(path:str, 
                threshold:int=50,
                force_close:int=6, 
                allow_merge:int=7, 
                sent_separator='<\>', 
                meta_header='', 
                **kwargs) -> Tuple[List[str], List[str]]:
    document = []
    if is_scanned_pdf(path, threshold):
        document = ocr_extract_text(path, **kwargs)
    else:
        _document = textract.process(path, extension='pdf', method='pdfminer')
        _document = [p for p in _document.decode('utf-8').split('\n') if p.strip() != '']
        document = requests.post(url=f'http://localhost:{kwargs["port"]}/', json={"document": _document}).json().get("generated_text")
        # document = _document
        del _document
    # print(document)
    if len(document) == 0:
        document = ocr_extract_text(path, **kwargs)
    pdf_declaration = []
    while len(document) > 0 and get_tag(document[0]) not in [1,2,3,4,5]:
        text = document.pop(0).strip('\n').strip().replace('\n', '')
        if len(text) == 0: continue
        if text[0].islower() and len(pdf_declaration) > 0:
            pdf_declaration[-1] += ' ' + text
        pdf_declaration.append(text)
    
    head_list, doc_list = extract_docx_layout(document, force_close, allow_merge, sent_separator, meta_header)
    head_list.insert(0, 'THÔNG TIN CHUNG')
    doc_list.insert(0, sent_separator.join(pdf_declaration))
    del document
    return head_list, doc_list
    

# %% [markdown]
# ---

# %% [markdown]
# ## Hậu xử lý

# %% [markdown]
# ### Load tokenizer

# %%
from transformers import AutoTokenizer

# %% [markdown]
# ### Concat các đoạn quá ngắn và cắt nhỏ các đoạn quá dài

# %% [markdown]
# #### Truncate các đoạn quá dài

# %%
def truncate(content:str, 
             tokenizer, 
             add_special_tokens:bool=True,
             max_tokens:int=300, 
             overlapse:float=0.3,
             sent_separator='<\>',
             alternative_sent_separator=['.', ';', '?', '!', '\n', '\t', ' ']) -> List[str]:
    if overlapse < 0: overlapse = 0
    elif overlapse > 0.5: overlapse = 0.5
    content_split = content.split(sent_separator)
    # print(content_split)
    token_count = [len(tl) \
        for tl in tokenizer(content_split, add_special_tokens=add_special_tokens).input_ids]
    
    result = []
    cache_content = []
    cache_count = []
    for sent, count in zip(content_split, token_count):
        if count > max_tokens:
            try:
                result += truncate(sent, tokenizer, add_special_tokens, max_tokens, overlapse, 
                                alternative_sent_separator.pop(0), alternative_sent_separator)
            except:
                result += truncate(sent[:len(sent)//2], tokenizer, add_special_tokens, max_tokens, overlapse)
                result += truncate(sent[len(sent)//2:], tokenizer, add_special_tokens, max_tokens, overlapse)
        elif sum(cache_count) + count > max_tokens and len(cache_content) != 0:
            result.append(sent_separator.join(cache_content))
            if overlapse > 0:
                retain_pos = int(len(cache_content) * (1-overlapse))
                cache_content = cache_content[retain_pos:]
                cache_count = cache_count[retain_pos:]
            else:
                cache_content = []
                cache_count = []
        cache_content.append(sent)
        cache_count.append(count)
    #end-for
    if len(cache_content) != 0:
        result.append(sent_separator.join(cache_content))
    del token_count
    return result
    
# %% [markdown]
# #### Merge các đoạn ngắn vào đoạn sau nó và truncate

# %%
def post_process(head_list:List[str], 
           doc_list:List[str], 
           tokenizer, 
           add_special_tokens:bool=False,
           max_tokens:int=300,
           overlapse:float=0.2,
           sent_separator='<\>') -> Tuple[List[str], List[str]]:
    
    assert len(head_list) == len(doc_list), f'Got unexpected document\'s layout!'
    doc_list_segment = []
    for _doc in doc_list:
        response = requests.post(url='http://localhost:9091/segment2', 
                                 json={'text': _doc.replace(sent_separator, ' '), 'sent_separator': sent_separator})
        doc_list_segment.append(response.json().get('sent'))
    tokens_count = [len(tl) \
        for tl in tokenizer([d.replace(sent_separator, ' ') for d in doc_list_segment], 
                            add_special_tokens=add_special_tokens).input_ids]
    # print(tokens_count)
    i = 0
    while i+1 < len(head_list):
        _head = head_list[i].strip()
        if tokens_count[i] + tokens_count[i+1] < max_tokens and _head == head_list[i+1].strip():
            _merge_head = head_list.pop(i+1)
            _merge_doc = doc_list_segment.pop(i+1)
            _merge_count = tokens_count.pop(i+1)
            # if _merge_head.strip() != _head.strip() and _merge_head.strip() != '':
            #     _merge_head = sent_separator.join([
            #                 h for h in _merge_head.split(sent_separator) \
            #                     if h not in _head.strip().split(sent_separator) \
            #                             and h.strip().__len__() > 0
            #                                 ])
                # head_list[i] += sent_separator + 'Và' + sent_separator + _merge_head
                # head_list[i] += sent_separator + _merge_head
            tokens_count[i] += _merge_count
            doc_list_segment[i] += sent_separator + _merge_doc
        else:
            i += 1 
    #end-while
        
    assert len(head_list) == len(doc_list_segment), f'Got unexpected document\'s layout!'
    
    result_head = []
    result_doc = []
    for _head, _doc, _count in zip(head_list, doc_list_segment, tokens_count):
        if _count > max_tokens:
            chunks = truncate(_doc, tokenizer, add_special_tokens, max_tokens, overlapse, sent_separator)
            for chunk in chunks:
                result_head.append(_head)
                result_doc.append(chunk)
        else:
            result_head.append(_head)
            result_doc.append(_doc)
            
    assert len(result_head) == len(result_doc), f'Got unexpected document\'s layout!'   
    del tokens_count
    return result_head, result_doc
            
# %% [markdown]
# ---

# %% [markdown]
# ## Hình thành pipeline

# %%
def get_attm_info_form(
    page_url=None,
    title=None,
    category=None,
    symbol_number=None,
    field=None,
    description=None,
    attachment_url=None,
    header=None,
    content=None
):
    return {
        "page_url"      : page_url,
        "title"         : title,
        "category"      : category,
        "symbol_number" : symbol_number,
        "field"         : field,
        "description"   : description,
        # "attachment_url": attachment_url,
        "header"        : header,
        "content"       : content
    }

# %%
def main(opt:Dict[str, str], global_opt:Dict[str, str]):
    for item in json.load(open(opt["attm_info_path"], 'r')):
        if "attachment_url" not in item.keys():
            item["attachment_url"] = None
        
        item_id = item["item_id"]
        item_attm_path = []
        for ex in global_opt["extensions"]:
            item_attm_path += [os.path.join(opt["attm_path"], ex, fp) \
                    for fp in os.listdir(os.path.join(opt["attm_path"], ex)) \
                        if fp.startswith(f"{item_id}-")]
        
        # print(item_id, ':', item_attm_path)
        form = get_attm_info_form(
                page_url        =   item["page_url"],
                title           =   item["title"],
                category        =   item["meta_data"]["Loại văn bản"],
                symbol_number   =   item["meta_data"]["Số ký hiệu"],
                field           =   item["meta_data"]["Lĩnh vực / Loại hình công việc"],
                description     =   f"""Ngày ban hành: {item["meta_data"]["Ngày ban hành"]}\n""" \
                                    f"""Ngày có hiệu lực: {item["meta_data"]["Ngày có hiệu lực"]}\n""" \
                                    f"""Tình trạng hiệu lực: {item["meta_data"]["Tình trạng hiệu lực"]}\n""" \
                                    f"""Ngày hết hiệu lực: {item["meta_data"]["Ngày hết hiệu lực"]}\n""" \
                                    f"""Cơ quan ban hành/ Người ký: {item["meta_data"]["Cơ quan ban hành/ Người ký"]}\n""",
                attachment_url  = item["attachment_url"],
            ) if opt["category"] == "document" else \
               get_attm_info_form(
                page_url        =   item["page_url"],
                title           =   item["title"],
                category        =   item["rule_category"],
                symbol_number   =   item["symbol_number"],
                field           =   item["field"],
                description     =   item["description"],
                attachment_url  =   item["attachment_url"],
            )

        for fp in item_attm_path:
            print(fp)
            if fp.endswith('.docx'):
                head_list, doc_list = extract_docx_layout(docx.Document(fp), 
                                                    force_close=opt["force_close"], 
                                                    allow_merge=opt["allow_merge"], 
                                                    sent_separator=global_opt["sent_separator"])
                
                post_headers, post_docs = post_process(head_list, doc_list, 
                                                    global_opt["tokenizer"], 
                                                    add_special_tokens=global_opt["add_special_tokens"],
                                                    max_tokens=global_opt["max_tokens"],
                                                    overlapse=global_opt["overlapse"],
                                                    sent_separator=global_opt["sent_separator"])
                
                for _head, _doc in zip(post_headers, post_docs):
                    form["header"] = _head
                    form["content"] = _doc
                    yield form
                    # print(_head)
                    # print('---')
                    # print(_doc)
                    # print('===============')
            elif fp.endswith('.pdf'):
                head_list, doc_list = extract_pdf_layout(fp,
                                                    threshold=global_opt["pdf_confidence_threshold"],
                                                    force_close=opt["force_close"], 
                                                    allow_merge=opt["allow_merge"], 
                                                    sent_separator=global_opt["sent_separator"])
                
                post_headers, post_docs = post_process(head_list, doc_list, 
                                                    global_opt["tokenizer"], 
                                                    add_special_tokens=global_opt["add_special_tokens"],
                                                    max_tokens=global_opt["max_tokens"],
                                                    overlapse=global_opt["overlapse"],
                                                    sent_separator=global_opt["sent_separator"])
                
                for _head, _doc in zip(post_headers, post_docs):
                    form["header"] = _head
                    form["content"] = _doc
                    yield form
                    # print(form)
                    # print(_head)
                    # print('---')
                    # print(_doc)
                    # print('===============')


# %%
# json_data = [item for item in main(doc_config, global_config)]

# %% [markdown]
# ---

# %% [markdown]
# ## Multiprocessing

# %%
import multiprocessing as mpc

# %%
def get_curr_timestamp() -> str:
    current_time = current_time = datetime.now()
    return current_time.strftime("%y/%m/%d %H:%M:%S")

def clean_data(text:str) -> str:
    while '  ' in text: text = text.replace('  ', ' ')
    while 'Và Và ' in text: text = text.replace('Và Và ', 'Và ')
    text = text.replace('Và \"', "\"")
    return text
# %%
def extract_layout_multiprocess(inputs:mpc.Queue, out_dir:str, index:int, closed_processors):
    # corrector = mpc.Process(target=call_flask_server, args=(MAGIC_PORT - index,))
    # corrector.start()
    while True:
        if inputs.empty(): continue
        task = inputs.get()
        if task is None: 
            with closed_processors.get_lock():
                closed_processors.value += 1
            print(f'{get_curr_timestamp()} PROCESS {index} | Close!')
            # corrector.terminate()
            break
        fp, opt, global_opt, form = task
        try:
            print(f'{get_curr_timestamp()} PROCESS {index} | Start extracting \"{fp}\" ...')
            
            global_opt["ocr_kwargs"]["port"] = MAGIC_PORTS[index % len(MAGIC_PORTS)]
            
            head_list, doc_list = extract_docx_layout(docx.Document(fp), 
                                                force_close=opt["force_close"], 
                                                allow_merge=opt["allow_merge"], 
                                                sent_separator=global_opt["sent_separator"]) \
                                        if fp.endswith('.docx') else \
                                    extract_pdf_layout(fp,
                                                threshold=global_opt["pdf_confidence_threshold"],
                                                force_close=opt["force_close"], 
                                                allow_merge=opt["allow_merge"], 
                                                sent_separator=global_opt["sent_separator"], 
                                                meta_header='',
                                                **global_opt["ocr_kwargs"])
            print(f'{get_curr_timestamp()} PROCESS {index} | Got {len(head_list)} blocks!')
            
            for overlapse in global_opt["overlapse"]:
                post_headers, post_docs = post_process(head_list.copy(), doc_list.copy(), 
                                                            global_opt["tokenizer"], 
                                                            add_special_tokens=global_opt["add_special_tokens"],
                                                            max_tokens=global_opt["max_tokens"],
                                                            overlapse=overlapse,
                                                            sent_separator=global_opt["sent_separator"])
                save_dir = os.path.join(out_dir, f'overlapse_{int(overlapse * 100):2d}')
                if not os.path.exists(save_dir):
                    os.mkdir(save_dir)
                print(f'{get_curr_timestamp()} PROCESS {index} | Post-processed got {len(post_headers)} chunks!')
                f = open(os.path.join(save_dir, f'proc-{index}.txt'), 'a')
                chunk_index = 0
                for _head, _doc in zip(post_headers, post_docs):
                    form["header"] = clean_data(_head.replace(global_opt["sent_separator"], ' '))
                    form["content"] = _doc.replace(global_opt["sent_separator"], ' ').replace('_', ' ')
                    form["chunk_index"] = chunk_index
                    chunk_index += 1
                    # chunks.append(form)
                    print(json.dumps(form, ensure_ascii=False), file=f, end='\n')
                del post_headers
                del post_docs
                print(f'{get_curr_timestamp()} PROCESS {index} | Save to {f.name}!')
                f.close()
            del head_list
            del doc_list
                
        except Exception as e: 
            error_path = os.path.join(out_dir, 'error')
            if not os.path.exists(error_path):
                os.mkdir(error_path)
            f = open(os.path.join(error_path, f'proc-{index}.txt'), 'a')
            form["error"] = e.args
            print(json.dumps(form, ensure_ascii=False), file=f, end='\n')
            f.close()
            print(f'{get_curr_timestamp()} PROCESS {index} | Got error! Save to {f.name}!')        
        
from tqdm import tqdm
# %%
def main_multiprocess(opt:Dict[str, str], global_opt:Dict[str, str], fp_queue:mpc.Queue, ignore_symbol_numbers:List[str]=[]):
    for item in tqdm(json.load(open(opt["attm_info_path"], 'r'))):
        if "attachment_url" not in item.keys():
            item["attachment_url"] = None
        
        form = get_attm_info_form(
                page_url        =   item["page_url"],
                title           =   item["title"],
                category        =   item["meta_data"]["Loại văn bản"],
                symbol_number   =   item["meta_data"]["Số ký hiệu"],
                field           =   item["meta_data"]["Lĩnh vực / Loại hình công việc"],
                description     =   f"""Ngày ban hành: {item["meta_data"]["Ngày ban hành"]}; """ \
                                    f"""Ngày có hiệu lực: {item["meta_data"]["Ngày có hiệu lực"]}; """ \
                                    f"""Tình trạng hiệu lực: {item["meta_data"]["Tình trạng hiệu lực"]}; """ \
                                    f"""Ngày hết hiệu lực: {item["meta_data"]["Ngày hết hiệu lực"]}; """ \
                                    f"""Cơ quan ban hành/ Người ký: {item["meta_data"]["Cơ quan ban hành/ Người ký"]};""",
                # attachment_url  = item["attachment_url"],
            ) if opt["category"] == "document" else \
               get_attm_info_form(
                page_url        =   item["page_url"],
                title           =   item["title"],
                category        =   item["rule_category"],
                symbol_number   =   item["symbol_number"],
                field           =   item["field"],
                description     =   item["description"],
                # attachment_url  =   item["attachment_url"],
            )
        if form["symbol_number"] in ignore_symbol_numbers: continue
        
        item_id = item["item_id"]
        item_attm_path = []
        for ex in global_opt["extensions"]:
            item_attm_path += [os.path.join(opt["attm_path"], ex, fp) \
                    for fp in os.listdir(os.path.join(opt["attm_path"], ex)) \
                        if fp.startswith(f"{item_id}-")]

        for fp in item_attm_path:  
            fp_queue.put([fp, opt, global_opt, form])

#%%
import time
def calculate_time(input:mpc.Value, num_proc:int, queue:mpc.Queue, step=10):
    start = time.time()
    qsize = queue.qsize()
    while input.value < num_proc and queue.qsize() > 0:
        execute_time = time.time()-start
        hh = execute_time // 3600
        mm = (execute_time % 3600) // 60
        ss = (execute_time % 3600) % 60
        print(f"""{get_curr_timestamp()} GLOBAL    | Execute time: {int(hh):02d}h-{int(mm):02d}m-{int(ss):02d}s; """ \
              f"""Progress: {qsize - queue.qsize()}/{qsize} ({100 - queue.qsize() * 100 // qsize}%)""")
        time.sleep(step)
    print(f"{get_curr_timestamp()} GLOBAL | STOP!!!")
        
#%%
import subprocess
def wakeup_corrector(host:str, port:int):
    # flask --app script_corrector.py run --host=0.0.0.0 --port=9298
    return subprocess.call(
        args=["flask", "--app=script_corrector.py", "run", f"--host={host}", f"--port={port}"],
        # env={"PATH": "/workspace/nlplab/kienvt/KLTN/kenv/bin"}, # /workspace/nlplab/kienvt/KLTN/kenv/bin
        # stderr=subprocess.DEVNULL,
        stdout=None,
    )
#%%

#%%
def get_ckpt_file_path(path) -> List[str]:
    results = []
    for _p in os.listdir(path):
        p = os.path.join(path, _p)
        if os.path.isfile(p):
            results.append(p)
        elif os.path.isdir(p):
            results += get_ckpt_file_path(p)
    return results
            
#%%

from argparse import ArgumentParser
from easydict import EasyDict
import yaml

if __name__ == "__main__":
    parser = ArgumentParser()    
    parser.add_argument("--config", help="Path to config file", type=str, default='./config/default.yml')
    opt = parser.parse_args()
    
    with open(opt.config, 'r') as f:
        conf = yaml.safe_load(f)
        f.close()
        
    args = EasyDict(conf)  
    del conf 
    
    global_config = args
    global_config.tokenizer = AutoTokenizer.from_pretrained(args.tokenizer, cache=args.hf_cache) 
    global_config.ocr_kwargs.margin = tuple(global_config.ocr_kwargs.margin)
    global_config.ocr_kwargs.im_size = tuple(global_config.ocr_kwargs.im_size)
    
    if not os.path.exists(args.out_dir):
        os.mkdir(args.out_dir)
    
    ignore_symbol_numbers = []
    if args.checkpoint_dir is not None and os.path.exists(args.checkpoint_dir):
        ckpt_path = get_ckpt_file_path(args.checkpoint_dir)
        for fp in ckpt_path:
            lines = open(fp, 'r').readlines()
            for line in lines:
                try:
                    item = json.loads(line)
                    ignore_symbol_numbers.append(item["symbol_number"])
                except:
                    pass
        ignore_symbol_numbers = list(set(ignore_symbol_numbers))

    
    NUM_POST_PROCESSOR = args.num_worker
    global MAGIC_PORTS
    MAGIC_PORTS = args.corrector_ports # list(range(MAGIC_PORT+1 - args.num_corrector, MAGIC_PORT+1))
    
    fp_queue = mpc.Queue()
    
    for item in args.attm_config:
        attm_type = item.keys()
        assert len(attm_type) == 1, f'Got unexpected config parameters!'
        attm_type = list(attm_type)[0]
        attm_conf = item.get(attm_type)
        main_multiprocess(attm_conf, global_config, fp_queue, ignore_symbol_numbers)
    
    # main_multiprocess(doc_config, global_config, fp_queue, ignore_symbol_numbers)
    # main_multiprocess(rule_config, global_config, fp_queue, ignore_symbol_numbers)
    
    for i in range(NUM_POST_PROCESSOR):
        fp_queue.put(None)
        
    correctors = []
    for i in range(args.num_corrector):
        corrector = mpc.Process(target=wakeup_corrector, 
                                args=("0.0.0.0", MAGIC_PORTS[i]), 
                                daemon=True)
        correctors.append(corrector)
        corrector.start()
    
    while True:
        try:
            for port in MAGIC_PORTS:
                response = requests.get(url=f'http://localhost:{port}/health')
                if response.status_code == 200: pass
                else: raise Exception("Error")
            break
        except:
            pass
    
    processors = []
    closed_processors = mpc.Value('d', 0)
    for i in range(NUM_POST_PROCESSOR):
        processor = mpc.Process(target=extract_layout_multiprocess, 
                                args=(fp_queue, args.out_dir, i, closed_processors), 
                                daemon=True)
        processors.append(processor)
        processor.start()
    
    global_processor = mpc.Process(target=calculate_time, 
                                    args=(closed_processors, NUM_POST_PROCESSOR, fp_queue, 300), 
                                    daemon=True)
    global_processor.start()
    
    while True:
        if closed_processors.value == NUM_POST_PROCESSOR:
            global_processor.join(timeout=1.0)
            del global_processor
            for processor in processors:
                processor.join(timeout=1.0)
                del processor
            for corrector in correctors:
                corrector.join(timeout=1.0)  
                del corrector
                # if not corrector.is_alive(): corrector.join(timeout=1.0)           
            fp_queue.close()
            break